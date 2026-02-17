"""
Resume Parser Utility
Extracts raw text from PDF and DOCX resume files.
Uses pypdf (lightweight, no Pillow dependency) instead of pdfplumber.
"""

import io
import os
import re
import logging

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF using pypdf (tiny, no Pillow needed)."""
    # Try pypdf first (new lightweight library)
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        parts  = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        text = clean_text("\n".join(parts))
        if text and len(text.strip()) > 30:
            return text
    except ImportError:
        pass
    except Exception as exc:
        logger.warning("pypdf failed: %s", exc)

    # Fallback: PyPDF2 (also installed)
    try:
        import PyPDF2
        parts = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return clean_text("\n".join(parts))
    except Exception as exc:
        logger.warning("PyPDF2 failed: %s", exc)

    raise RuntimeError("Could not extract text from PDF. Try a different file.")


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc  = Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return clean_text("\n".join(parts))
    except Exception as exc:
        logger.warning("python-docx failed: %s", exc)
        raise RuntimeError(f"Could not parse DOCX: {exc}")


def extract_text_from_file(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
